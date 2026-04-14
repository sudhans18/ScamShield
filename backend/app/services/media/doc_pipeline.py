"""
Document extraction + forgery heuristics for ScamShield.
"""

from __future__ import annotations

import io
import logging
import re
from pathlib import Path
from typing import Optional, Union

logger = logging.getLogger(__name__)

MAX_DOC_BYTES = 20 * 1024 * 1024
MAX_PAGES_OCR = 10
SUPPORTED_FORMATS = {".pdf", ".docx", ".doc"}

KNOWN_COMPANIES = [
    "Tata Projects",
    "Tata Consultancy Services",
    "Tata Motors",
    "Infosys",
    "Wipro",
    "HCL Technologies",
    "Tech Mahindra",
    "Larsen and Toubro",
    "L&T",
    "Reliance Industries",
    "HDFC Bank",
    "ICICI Bank",
    "State Bank of India",
    "Axis Bank",
    "Hindustan Unilever",
    "ITC Limited",
    "Maruti Suzuki",
    "Mahindra and Mahindra",
    "Bajaj Auto",
    "Hero MotoCorp",
    "BHEL",
    "ONGC",
    "NTPC",
    "Coal India",
    "SAIL",
    "GAIL",
]

COMPANY_SUFFIX = (
    r"(?:Pvt\.?\s*Ltd\.?|Private\s+Limited|Ltd\.?|Limited|Inc\.?|LLP|"
    r"Enterprises?|Solutions?|Services?|Consultanc(?:y|ies)|Agenc(?:y|ies)|"
    r"Recruitment|International|Intl\.?|Overseas|Manpower|Technologies|"
    r"Tech\.?|Group|Corp\.?|Corporation)"
)
COMPANY_RE = re.compile(rf"([A-Z][A-Za-z0-9\s&.\-]{{2,45}}?)\s*{COMPANY_SUFFIX}", re.IGNORECASE)
GST_RE = re.compile(r"\b(\d{2}[A-Z]{5}\d{4}[A-Z][1-9A-Z]Z[0-9A-Z])\b")
CIN_RE = re.compile(r"\b([LU]\d{5}[A-Z]{2}\d{4}(?:PLC|PTC|OPC|NPL|GAP|GOI)\d{6})\b")
FEE_IN_DOC_RE = re.compile(
    r"(?:registration|medical|processing|security|visa|advance|joining)"
    r"\s*(?:fee|fees|charge|deposit)",
    re.IGNORECASE,
)
URGENCY_IN_DOC_RE = re.compile(
    r"\b(?:urgent|immediately|asap|today only|limited time|act now)\b",
    re.IGNORECASE,
)


def process_document(source: Union[str, bytes, Path], filename: str = "") -> dict:
    """
    Extract text from a PDF/DOCX and run forgery heuristics.
    """
    load = _load(source, filename)
    if not load["success"]:
        return load

    raw_bytes = load["raw_bytes"]
    doc_format = load["doc_format"]

    extract = _extract_text(raw_bytes, doc_format)
    if not extract["success"]:
        return {**extract, "doc_format": doc_format}

    text = extract["text"]
    method = extract["method"]
    pages = extract.get("page_count", 1)
    if not text.strip():
        return _err(
            "No readable text found. Document may be a blank scan or image-only PDF.",
            extra={"doc_format": doc_format, "extraction_method": method},
        )

    company_name = _extract_company(text)
    gst_number = _extract_gst(text)
    cin_number = _extract_cin(text)
    typo = _typosquatting_check(company_name) if company_name else {"suspicious": False}
    forgery = _score_forgery(
        text=text,
        company=company_name,
        gst=gst_number,
        typosquatting_suspected=bool(typo.get("suspicious")),
    )

    logger.info(
        "doc_pipeline: format=%s method=%s pages=%s company=%r gst=%r typo=%s risk=%s",
        doc_format,
        method,
        pages,
        company_name,
        gst_number,
        bool(typo.get("suspicious")),
        forgery["risk"],
    )

    return {
        "success": True,
        "extracted_text": text,
        "doc_format": doc_format,
        "page_count": pages,
        "extraction_method": method,
        "company_name": company_name,
        "gst_number": gst_number,
        "cin_number": cin_number,
        "typosquatting_detected": bool(typo.get("suspicious")),
        "typosquatting_similar_to": typo.get("similar_to"),
        "forgery_risk": forgery["risk"],
        "forgery_reasons": forgery["reasons"],
        "char_count": len(text),
    }


def _load(source: Union[str, bytes, Path], filename: str) -> dict:
    try:
        if isinstance(source, (str, Path)):
            path = Path(source)
            if not path.exists():
                return _err(f"File not found: {path}")
            if path.stat().st_size > MAX_DOC_BYTES:
                return _err(f"File too large ({path.stat().st_size // 1024 // 1024} MB, max 20 MB).")
            ext = path.suffix.lower()
            if ext not in SUPPORTED_FORMATS:
                return _err(f"Unsupported format '{ext}'. Supported: {SUPPORTED_FORMATS}")
            with open(path, "rb") as file_obj:
                raw = file_obj.read()
            return {"success": True, "raw_bytes": raw, "doc_format": _detect_format(ext, raw)}

        if isinstance(source, (bytes, bytearray)):
            raw = bytes(source)
            if len(raw) > MAX_DOC_BYTES:
                return _err("Document too large (max 20 MB).")
            ext = Path(filename).suffix.lower() if filename else ""
            fmt = _detect_format(ext, raw)
            if fmt == "unknown":
                return _err(f"Cannot determine document format from extension '{ext}'.")
            return {"success": True, "raw_bytes": raw, "doc_format": fmt}

        return _err(f"Invalid source type: {type(source).__name__}")
    except Exception as exc:
        logger.exception("doc_pipeline: _load failed")
        return _err(f"Could not load document: {exc}")


def _detect_format(ext: str, raw: bytes) -> str:
    if ext == ".pdf" or raw[:4] == b"%PDF":
        return "pdf"
    if ext in (".docx", ".doc") or raw[:2] == b"PK":
        return "docx"
    return "unknown"


def _extract_text(raw: bytes, fmt: str) -> dict:
    if fmt == "pdf":
        return _extract_pdf(raw)
    if fmt == "docx":
        return _extract_docx(raw)
    return {"success": False, "error": f"No extractor for format: {fmt}"}


def _extract_pdf(raw: bytes) -> dict:
    import pdfplumber

    pages_text: list[str] = []
    ocr_pages: list[tuple[int, object]] = []
    page_count = 0

    try:
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            if pdf.metadata and pdf.metadata.get("Encrypted"):
                return {"success": False, "error": "PDF is encrypted/password-protected. Cannot extract text."}

            page_count = len(pdf.pages)
            for index, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        pages_text.append(page_text)
                    else:
                        ocr_pages.append((index, page))
                except Exception:
                    ocr_pages.append((index, page))
    except pdfplumber.PDFSyntaxError as exc:
        return {"success": False, "error": f"Corrupted or invalid PDF: {exc}"}
    except Exception as exc:
        logger.warning("doc_pipeline: pdfplumber failed (%s), falling back to full OCR", exc)
        return _extract_pdf_ocr(raw)

    if not pages_text and ocr_pages:
        return _extract_pdf_ocr(raw)

    if ocr_pages:
        for index, page in ocr_pages[:MAX_PAGES_OCR]:
            ocr_text = _ocr_pdfplumber_page(page)
            if ocr_text:
                pages_text.insert(index, ocr_text)

    combined = "\n\n".join(pages_text)
    method = "mixed" if ocr_pages else "pdfplumber"
    return {"success": True, "text": combined, "method": method, "page_count": page_count}


def _ocr_pdfplumber_page(page: object) -> str:
    try:
        import cv2
        import numpy as np
        import pytesseract

        pil_img = page.to_image(resolution=200).original
        gray = cv2.cvtColor(np.array(pil_img.convert("RGB")), cv2.COLOR_RGB2GRAY)
        _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        text = pytesseract.image_to_string(binary, lang="eng+hin", config="--psm 6 --oem 3")
        return text.strip()
    except Exception as exc:
        logger.warning("doc_pipeline: page OCR failed - %s", exc)
        return ""


def _extract_pdf_ocr(raw: bytes) -> dict:
    try:
        from pdf2image import convert_from_bytes
    except ImportError:
        return {
            "success": False,
            "error": (
                "PDF appears scanned but pdf2image is not installed. "
                "Install: pip install pdf2image and install poppler."
            ),
        }

    try:
        import cv2
        import numpy as np
        import pytesseract

        images = convert_from_bytes(raw, dpi=200, first_page=1, last_page=MAX_PAGES_OCR)
        all_text: list[str] = []
        for pil_img in images:
            gray = cv2.cvtColor(np.array(pil_img.convert("RGB")), cv2.COLOR_RGB2GRAY)
            _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            text = pytesseract.image_to_string(binary, lang="eng+hin", config="--psm 6 --oem 3")
            all_text.append(text.strip())

        combined = "\n\n".join(filter(None, all_text))
        return {"success": True, "text": combined, "method": "ocr", "page_count": len(images)}
    except Exception as exc:
        logger.exception("doc_pipeline: full PDF OCR failed")
        return {"success": False, "error": f"PDF OCR failed: {exc}"}


def _extract_docx(raw: bytes) -> dict:
    try:
        import docx

        doc = docx.Document(io.BytesIO(raw))
        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        for section in doc.sections:
            for hdr_para in section.header.paragraphs:
                text = hdr_para.text.strip()
                if text:
                    parts.insert(0, f"[HEADER] {text}")

        return {"success": True, "text": "\n".join(parts), "method": "python-docx", "page_count": 1}
    except ImportError:
        return {"success": False, "error": "python-docx not installed. Run: pip install python-docx"}
    except Exception as exc:
        if "not a valid" in str(exc).lower() or "invalid" in str(exc).lower():
            return {
                "success": False,
                "error": "File may be legacy .doc format. Ask sender to save as .docx or PDF.",
            }
        return {"success": False, "error": f"DOCX extraction failed: {exc}"}


def _extract_company(text: str) -> Optional[str]:
    matches = COMPANY_RE.findall(text)
    if not matches:
        return None
    return " ".join(matches[0].split())


def _extract_gst(text: str) -> Optional[str]:
    match = GST_RE.search(text)
    return match.group(1) if match else None


def _extract_cin(text: str) -> Optional[str]:
    match = CIN_RE.search(text)
    return match.group(1) if match else None


def _valid_gst_format(gst: str) -> bool:
    return bool(re.match(r"^\d{2}[A-Z]{5}\d{4}[A-Z][1-9A-Z]Z[0-9A-Z]$", gst))


def _typosquatting_check(company: str) -> dict:
    company_lower = company.lower().strip()
    for known in KNOWN_COMPANIES:
        distance = _levenshtein(company_lower, known.lower())
        if 0 < distance <= 3:
            return {"suspicious": True, "similar_to": known, "distance": distance}
    return {"suspicious": False, "similar_to": None, "distance": None}


def _levenshtein(a: str, b: str) -> int:
    if len(a) < len(b):
        return _levenshtein(b, a)
    if not b:
        return len(a)
    prev = list(range(len(b) + 1))
    for index_a, char_a in enumerate(a):
        curr = [index_a + 1]
        for index_b, char_b in enumerate(b):
            curr.append(min(prev[index_b + 1] + 1, curr[index_b] + 1, prev[index_b] + (char_a != char_b)))
        prev = curr
    return prev[-1]


def _score_forgery(
    text: str,
    company: Optional[str],
    gst: Optional[str],
    typosquatting_suspected: bool,
) -> dict:
    score = 0
    reasons: list[str] = []

    if not company:
        score += 25
        reasons.append("No company name found in document")

    if not gst:
        score += 20
        reasons.append("No GST number found")
    elif not _valid_gst_format(gst):
        score += 20
        reasons.append("GST format appears invalid")

    if typosquatting_suspected:
        score += 25
        reasons.append("Company name resembles a known brand (possible impersonation)")

    if FEE_IN_DOC_RE.search(text):
        score += 40
        reasons.append("Document asks for upfront fee")

    if URGENCY_IN_DOC_RE.search(text):
        score += 15
        reasons.append("Urgency language is unusual for formal offer documents")

    score = min(score, 100)
    risk = "high" if score >= 60 else "medium" if score >= 30 else "low"
    return {"risk": risk, "score": score, "reasons": reasons[:3]}


def _err(message: str, extra: dict | None = None) -> dict:
    logger.warning("doc_pipeline: %s", message)
    payload = {"success": False, "error": message}
    if extra:
        payload.update(extra)
    return payload

