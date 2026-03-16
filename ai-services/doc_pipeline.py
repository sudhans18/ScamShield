"""
doc_pipeline.py
---------------
Production document text extraction + forgery detection for ScamShield.

Supports:
  DOCX  — python-docx (headers, body paragraphs, tables)
  PDF (text-based)   — pdfplumber (layout-aware, handles columns)
  PDF (scanned/image)— pytesseract OCR per page via PIL/numpy
  PDF (encrypted)    — detects and reports gracefully
  PDF (mixed)        — pdfplumber first, OCR fallback per page

Forgery detection:
  - Company name extraction + Levenshtein typosquatting check
  - GST / CIN number extraction + format validation
  - Fee-in-offer-letter detection (definitive scam signal)
  - MCA21 cross-check (real HTTP call, graceful timeout)
  - Urgency language in formal documents

All errors return structured dicts — this function never raises.
"""

from __future__ import annotations

import io
import logging
import re
import tempfile
import os
from pathlib import Path
from typing import Optional, Union

logger = logging.getLogger(__name__)

MAX_DOC_BYTES     = 20 * 1024 * 1024   # 20 MB
MAX_PAGES_OCR     = 10                  # OCR only the first N pages of large scanned PDFs
API_TIMEOUT_S     = 5
SUPPORTED_FORMATS = {".pdf", ".docx", ".doc"}

KNOWN_COMPANIES = [
    "Tata Projects", "Tata Consultancy Services", "Tata Motors",
    "Infosys", "Wipro", "HCL Technologies", "Tech Mahindra",
    "Larsen and Toubro", "L&T", "Reliance Industries",
    "HDFC Bank", "ICICI Bank", "State Bank of India", "Axis Bank",
    "Hindustan Unilever", "ITC Limited", "Maruti Suzuki",
    "Mahindra and Mahindra", "Bajaj Auto", "Hero MotoCorp",
    "BHEL", "ONGC", "NTPC", "Coal India", "SAIL", "GAIL",
]

COMPANY_SUFFIX = (
    r"(?:Pvt\.?\s*Ltd\.?|Private\s+Limited|Ltd\.?|Limited|Inc\.?|LLP|"
    r"Enterprises?|Solutions?|Services?|Consultanc(?:y|ies)|Agenc(?:y|ies)|"
    r"Recruitment|International|Intl\.?|Overseas|Manpower|Technologies|"
    r"Tech\.?|Group|Corp\.?|Corporation)"
)
COMPANY_RE = re.compile(rf"([A-Z][A-Za-z0-9\s&.\-]{{2,45}}?)\s*{COMPANY_SUFFIX}", re.IGNORECASE)

GST_RE = re.compile(r"\b(\d{2}[A-Z]{5}\d{4}[A-Z][1-9A-Z]Z[0-9A-Z])\b")
CIN_RE = re.compile(r"\b([LU]\d{5}[A-Z]{2}\d{4}(?:PLC|PTC|OPC|NPL|GAP|GOI)\d{6})\b")
FEE_IN_DOC_RE = re.compile(
    r"(?:registration|medical|processing|security|visa|advance|joining)"
    r"\s*(?:fee|fees|charge|deposit)",
    re.IGNORECASE,
)
URGENCY_IN_DOC_RE = re.compile(
    r"\b(?:urgent|immediately|asap|today only|limited time|act now)\b",
    re.IGNORECASE,
)


# ── Public entry point ────────────────────────────────────────────────────────

def process_document(source: Union[str, bytes, Path], filename: str = "") -> dict:
    """
    Extract text from a document and run forgery analysis.

    Parameters
    ----------
    source   : str | bytes | Path
    filename : str — original filename (used to infer format from bytes)

    Returns
    -------
    dict:
        success          bool
        extracted_text   str
        doc_format       str   "pdf" | "docx" | "unknown"
        page_count       int
        extraction_method str  "pdfplumber" | "ocr" | "python-docx" | "mixed"
        company_name     str | None
        gst_number       str | None
        cin_number       str | None
        mca_verified     bool
        mca_message      str
        forgery_risk     str   "low" | "medium" | "high"
        forgery_reasons  list[str]
        char_count       int
        error            str   only on failure
    """
    # ── 1. Load ───────────────────────────────────────────────────────────────
    load = _load(source, filename)
    if not load["success"]:
        return load

    raw_bytes  = load["raw_bytes"]
    doc_format = load["doc_format"]

    # ── 2. Extract text ───────────────────────────────────────────────────────
    extract = _extract_text(raw_bytes, doc_format)
    if not extract["success"]:
        return {**extract, "doc_format": doc_format}

    text   = extract["text"]
    method = extract["method"]
    pages  = extract.get("page_count", 1)

    if not text.strip():
        return _err(
            "No readable text found. Document may be a blank scan or image-only PDF.",
            extra={"doc_format": doc_format, "extraction_method": method},
        )

    # ── 3. Entity extraction for forgery detection ────────────────────────────
    company_name = _extract_company(text)
    gst_number   = _extract_gst(text)
    cin_number   = _extract_cin(text)

    # ── 4. Registry cross-check ───────────────────────────────────────────────
    mca = _check_mca(company_name, gst_number, cin_number)

    # ── 5. Forgery scoring ────────────────────────────────────────────────────
    forgery = _score_forgery(text, company_name, gst_number, mca["verified"])

    logger.info(
        f"doc_pipeline: format={doc_format} method={method} pages={pages} "
        f"company={company_name!r} gst={gst_number!r} "
        f"mca={mca['verified']} forgery={forgery['risk']}"
    )

    return {
        "success":           True,
        "extracted_text":    text,
        "doc_format":        doc_format,
        "page_count":        pages,
        "extraction_method": method,
        "company_name":      company_name,
        "gst_number":        gst_number,
        "cin_number":        cin_number,
        "mca_verified":      mca["verified"],
        "mca_message":       mca["message"],
        "forgery_risk":      forgery["risk"],
        "forgery_reasons":   forgery["reasons"],
        "char_count":        len(text),
    }


# ── Loading ───────────────────────────────────────────────────────────────────

def _load(source: Union[str, bytes, Path], filename: str) -> dict:
    try:
        if isinstance(source, (str, Path)):
            path = Path(source)
            if not path.exists():
                return _err(f"File not found: {path}")
            if path.stat().st_size > MAX_DOC_BYTES:
                return _err(f"File too large ({path.stat().st_size // 1024 // 1024} MB, max 20 MB).")
            ext = path.suffix.lower()
            if ext not in SUPPORTED_FORMATS:
                return _err(f"Unsupported format '{ext}'. Supported: {SUPPORTED_FORMATS}")
            with open(path, "rb") as f:
                raw = f.read()
            return {"success": True, "raw_bytes": raw, "doc_format": _detect_format(ext, raw)}

        elif isinstance(source, (bytes, bytearray)):
            raw = bytes(source)
            if len(raw) > MAX_DOC_BYTES:
                return _err(f"Document too large (max 20 MB).")
            ext = Path(filename).suffix.lower() if filename else ""
            fmt = _detect_format(ext, raw)
            if fmt == "unknown":
                return _err(f"Cannot determine document format from extension '{ext}'.")
            return {"success": True, "raw_bytes": raw, "doc_format": fmt}

        else:
            return _err(f"Invalid source type: {type(source).__name__}")

    except Exception as e:
        logger.exception("doc_pipeline: _load failed")
        return _err(f"Could not load document: {e}")


def _detect_format(ext: str, raw: bytes) -> str:
    if ext == ".pdf" or raw[:4] == b"%PDF":
        return "pdf"
    if ext in (".docx", ".doc") or raw[:2] == b"PK":
        return "docx"
    return "unknown"


# ── Text extraction ───────────────────────────────────────────────────────────

def _extract_text(raw: bytes, fmt: str) -> dict:
    if fmt == "pdf":
        return _extract_pdf(raw)
    if fmt == "docx":
        return _extract_docx(raw)
    return {"success": False, "error": f"No extractor for format: {fmt}"}


def _extract_pdf(raw: bytes) -> dict:
    """
    PDF extraction with three-stage fallback:
    1. pdfplumber (layout-aware, best for structured documents)
    2. Per-page OCR for pages where pdfplumber returns nothing (scanned pages)
    3. Fully OCR-based for entirely scanned PDFs

    Handles encrypted PDFs gracefully.
    """
    import pdfplumber

    pages_text = []
    ocr_pages  = []
    page_count = 0

    try:
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            # Encrypted PDF detection
            if pdf.metadata and pdf.metadata.get("Encrypted"):
                return {"success": False, "error": "PDF is encrypted/password-protected. Cannot extract text."}

            page_count = len(pdf.pages)
            logger.info(f"doc_pipeline: PDF has {page_count} pages")

            for i, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        pages_text.append(page_text)
                    else:
                        # Page returned nothing — likely a scanned image
                        ocr_pages.append((i, page))
                except Exception as e:
                    logger.warning(f"doc_pipeline: pdfplumber error page {i+1} — {e}")
                    ocr_pages.append((i, page))

    except pdfplumber.PDFSyntaxError as e:
        return {"success": False, "error": f"Corrupted or invalid PDF: {e}"}
    except Exception as e:
        logger.warning(f"doc_pipeline: pdfplumber failed ({e}), falling back to full OCR")
        return _extract_pdf_ocr(raw)

    # If pdfplumber got nothing at all → full OCR
    if not pages_text and ocr_pages:
        logger.info("doc_pipeline: no text from pdfplumber, running full OCR")
        return _extract_pdf_ocr(raw)

    # Mixed: OCR the pages pdfplumber missed (up to MAX_PAGES_OCR)
    if ocr_pages:
        logger.info(f"doc_pipeline: {len(ocr_pages)} pages need OCR")
        for i, page in ocr_pages[:MAX_PAGES_OCR]:
            ocr_text = _ocr_pdfplumber_page(page)
            if ocr_text:
                pages_text.insert(i, ocr_text)

    combined = "\n\n".join(pages_text)
    method = "mixed" if ocr_pages else "pdfplumber"
    return {"success": True, "text": combined, "method": method, "page_count": page_count}


def _ocr_pdfplumber_page(page) -> str:
    """Rasterise a single pdfplumber page and OCR it."""
    try:
        import pytesseract
        from PIL import Image
        import numpy as np
        import cv2

        # Render page at 200 DPI as PIL image
        pil_img = page.to_image(resolution=200).original
        # Convert to greyscale and apply Otsu threshold
        gray = cv2.cvtColor(np.array(pil_img.convert("RGB")), cv2.COLOR_RGB2GRAY)
        _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        text = pytesseract.image_to_string(binary, lang="eng+hin", config="--psm 6 --oem 3")
        return text.strip()
    except Exception as e:
        logger.warning(f"doc_pipeline: page OCR failed — {e}")
        return ""


def _extract_pdf_ocr(raw: bytes) -> dict:
    """
    Full OCR fallback for entirely scanned PDFs.
    Converts each page to an image with pdf2image (if installed),
    then runs Tesseract. Falls back to error if pdf2image unavailable.
    """
    try:
        from pdf2image import convert_from_bytes
    except ImportError:
        return {
            "success": False,
            "error": (
                "PDF appears to be a scanned image but pdf2image is not installed. "
                "Install: pip install pdf2image  and  apt install poppler-utils"
            ),
        }

    try:
        import pytesseract
        import cv2
        import numpy as np

        images = convert_from_bytes(raw, dpi=200, first_page=1, last_page=MAX_PAGES_OCR)
        all_text = []
        for i, pil_img in enumerate(images):
            gray = cv2.cvtColor(np.array(pil_img.convert("RGB")), cv2.COLOR_RGB2GRAY)
            _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            text = pytesseract.image_to_string(binary, lang="eng+hin", config="--psm 6 --oem 3")
            all_text.append(text.strip())
            logger.info(f"doc_pipeline: OCR page {i+1}/{len(images)}")

        combined = "\n\n".join(filter(None, all_text))
        return {"success": True, "text": combined, "method": "ocr", "page_count": len(images)}

    except Exception as e:
        logger.exception("doc_pipeline: full PDF OCR failed")
        return {"success": False, "error": f"PDF OCR failed: {e}"}


def _extract_docx(raw: bytes) -> dict:
    """
    Extract text from DOCX including headers, body, and tables.
    python-docx only reads .docx (ZIP-based). Older .doc format
    is not supported — we report that clearly.
    """
    try:
        import docx
        doc = docx.Document(io.BytesIO(raw))

        parts = []

        # Body paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Tables (offer letters often use tables for salary/role details)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        # Headers and footers (company name often in header)
        for section in doc.sections:
            for hdr_para in section.header.paragraphs:
                text = hdr_para.text.strip()
                if text:
                    parts.insert(0, f"[HEADER] {text}")

        combined = "\n".join(parts)
        return {"success": True, "text": combined, "method": "python-docx", "page_count": 1}

    except ImportError:
        return {"success": False, "error": "python-docx not installed. Run: pip install python-docx"}
    except Exception as e:
        # Likely a .doc (legacy binary format) — report clearly
        if "not a valid" in str(e).lower() or "invalid" in str(e).lower():
            return {
                "success": False,
                "error": (
                    "File may be a legacy .doc format (not .docx). "
                    "Please ask the sender to save as .docx or convert to PDF."
                ),
            }
        return {"success": False, "error": f"DOCX extraction failed: {e}"}


# ── Forgery detection ─────────────────────────────────────────────────────────

def _extract_company(text: str) -> Optional[str]:
    matches = COMPANY_RE.findall(text)
    if not matches:
        return None
    return " ".join(matches[0].split())

def _extract_gst(text: str) -> Optional[str]:
    m = GST_RE.search(text)
    return m.group(1) if m else None

def _extract_cin(text: str) -> Optional[str]:
    m = CIN_RE.search(text)
    return m.group(1) if m else None

def _check_mca(company: Optional[str], gst: Optional[str], cin: Optional[str]) -> dict:
    """Cross-check against MCA21. Times out gracefully."""
    if not company and not gst and not cin:
        return {"verified": False, "message": "No company identifiers found to check."}

    if gst:
        if not _valid_gst_format(gst):
            return {"verified": False, "message": f"GST number '{gst}' has an invalid format."}

    if company:
        typo = _typosquatting_check(company)
        if typo["suspicious"]:
            return {
                "verified": False,
                "message": (
                    f"Company name '{company}' resembles '{typo['similar_to']}' "
                    f"(edit distance {typo['distance']}) — possible impersonation."
                ),
            }

    # Real MCA API attempt
    if company or cin:
        result = _mca_api_call(company, cin)
        if result:
            return result

    return {
        "verified": False,
        "message": (
            f"'{company or gst or cin}' could not be verified in MCA registry. "
            "Verify manually at https://www.mca.gov.in/mcafoportal/viewCompanyMasterData.do"
        ),
    }

def _valid_gst_format(gst: str) -> bool:
    """Structural GST format check (does not verify the checksum)."""
    return bool(re.match(r"^\d{2}[A-Z]{5}\d{4}[A-Z][1-9A-Z]Z[0-9A-Z]$", gst))

def _mca_api_call(company: Optional[str], cin: Optional[str]) -> Optional[dict]:
    """Attempt MCA21 lookup. Returns None on timeout/error."""
    try:
        import requests
        params = {}
        if cin:
            params["companyNumber"] = cin
        elif company:
            params["companyName"] = company

        resp = requests.get(
            "https://www.mca.gov.in/mcafoportal/viewCompanyMasterData.do",
            params=params, timeout=API_TIMEOUT_S,
        )
        if resp.status_code == 200 and len(resp.text) > 200:
            name_part = (company or "").split()[0].lower()
            if name_part and name_part in resp.text.lower():
                return {"verified": True, "message": "Company found in MCA registry."}
        return None
    except Exception as e:
        logger.warning(f"doc_pipeline: MCA API call failed — {e}")
        return None

def _typosquatting_check(company: str) -> dict:
    """Levenshtein distance check against known Indian companies."""
    company_lower = company.lower().strip()
    for known in KNOWN_COMPANIES:
        d = _levenshtein(company_lower, known.lower())
        if 0 < d <= 3:
            return {"suspicious": True, "similar_to": known, "distance": d}
    return {"suspicious": False}

def _levenshtein(a: str, b: str) -> int:
    if len(a) < len(b): return _levenshtein(b, a)
    if not b: return len(a)
    prev = list(range(len(b) + 1))
    for i, ca in enumerate(a):
        curr = [i + 1]
        for j, cb in enumerate(b):
            curr.append(min(prev[j+1]+1, curr[j]+1, prev[j]+(ca!=cb)))
        prev = curr
    return prev[-1]

def _score_forgery(text: str, company: Optional[str], gst: Optional[str], mca_verified: bool) -> dict:
    score   = 0
    reasons = []

    if not company:
        score += 25; reasons.append("No company name found in document")
    if not gst:
        score += 20; reasons.append("No GST number — all legitimate employers must have one")
    if company and not mca_verified:
        score += 20; reasons.append(f"'{company}' not verifiable in MCA registry")
    if FEE_IN_DOC_RE.search(text):
        score += 40; reasons.append("Document requests upfront fee — legitimate offers never do this")
    if URGENCY_IN_DOC_RE.search(text):
        score += 15; reasons.append("Urgency language in a formal offer letter is unusual")

    score = min(score, 100)
    risk = "high" if score >= 60 else "medium" if score >= 30 else "low"
    return {"risk": risk, "score": score, "reasons": reasons[:3]}


# ── Helpers ───────────────────────────────────────────────────────────────────

def _err(msg: str, extra: dict | None = None) -> dict:
    logger.warning(f"doc_pipeline: {msg}")
    base = {"success": False, "error": msg}
    if extra:
        base.update(extra)
    return base


# ── Manual test ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import json

    # Create a minimal test DOCX in memory
    try:
        import docx as _docx
        doc = _docx.Document()
        doc.add_heading("OFFER LETTER", 0)
        doc.add_paragraph("Company: Tata Projcts Ltd")
        doc.add_paragraph("GST: 27AABCT1234Z1ZQ")
        doc.add_paragraph("You are selected for Security Guard - Dubai.")
        doc.add_paragraph("Salary: Rs. 80,000 per month")
        doc.add_paragraph("IMPORTANT: Pay registration fee Rs. 8,000 immediately.")
        doc.add_paragraph("Contact: 9876543210 | Limited seats available.")
        buf = io.BytesIO()
        doc.save(buf)
        raw = buf.getvalue()
        print("Testing DOCX pipeline...\n")
        result = process_document(raw, filename="offer.docx")
        print(json.dumps(result, indent=2, ensure_ascii=False))
    except ImportError:
        print("python-docx not installed. Install: pip install python-docx")